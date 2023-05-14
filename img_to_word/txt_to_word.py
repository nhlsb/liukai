from docx import Document
import os

def extract_text_from_txt(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        text = file.read()
    return text

def merge_txt_to_word(txt_folder, output_file):
    document = Document()

    # 遍历TXT文件夹中的文件
    for filename in os.listdir(txt_folder):
        if filename.endswith(".txt"):
            txt_path = os.path.join(txt_folder, filename)
            text = extract_text_from_txt(txt_path)

            # 添加文本内容到Word文档中
            document.add_paragraph(text)

            # 添加换行分隔符
            document.add_paragraph("\n")

    # 保存合并后的Word文档
    document.save(output_file)

# 示例用法
txt_folder = "D:/img_to_word"
output_file = "D:/img_to_word/txt_to_word.docx"

merge_txt_to_word(txt_folder, output_file)
